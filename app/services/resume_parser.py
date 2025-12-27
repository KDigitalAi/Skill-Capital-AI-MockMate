"""
Resume parsing service using LangChain and PyMuPDF
Extracts skills and experience level from resume files
"""

import os
import tempfile
import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
import re

# Setup logger
logger = logging.getLogger(__name__)

# Try to import PyMuPDF, handle import errors and DLL errors gracefully
PYMUPDF_AVAILABLE = False
fitz = None
try:
    import fitz  # PyMuPDF
    # Test if it actually works (not just imported) - DLL might fail at runtime
    try:
        # Quick test: try to create an empty document
        test_doc = fitz.open()
        test_doc.close()
        PYMUPDF_AVAILABLE = True
    except Exception:
        # DLL error or other runtime error
        PYMUPDF_AVAILABLE = False
        fitz = None
except Exception as e:
    # Catch all exceptions including ImportError and DLL load errors on Windows
    PYMUPDF_AVAILABLE = False
    fitz = None

# Try to import python-docx, handle import errors gracefully
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

# Note: LangChain components are not currently used in this parser
# They are kept for potential future use

class ResumeParser:
    """Parse resume files and extract skills and experience"""
    
    def __init__(self):
        self.skill_keywords = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'php', 'ruby',
            'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'html', 'css', 'sass', 'less',
            # Frameworks & Libraries
            'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'fastapi',
            'spring', 'laravel', 'rails', 'asp.net', '.net', 'next.js', 'nuxt.js',
            # Databases
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
            'oracle', 'sqlite', 'dynamodb', 'firebase', 'supabase',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'ci/cd',
            'terraform', 'ansible', 'linux', 'bash', 'shell scripting',
            # Tools & Others
            'git', 'github', 'gitlab', 'jira', 'agile', 'scrum', 'rest api', 'graphql',
            'microservices', 'machine learning', 'ai', 'data science', 'nlp', 'computer vision'
        ]
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF with fallback"""
        if not os.path.exists(file_path):
            raise Exception(f"PDF file not found at path: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise Exception("PDF file is empty (0 bytes)")
        
        
        # Try PyMuPDF first if available
        if PYMUPDF_AVAILABLE:
            try:
                # Open PDF in binary mode
                doc = fitz.open(file_path)
                
                text = ""
                for page_num, page in enumerate(doc):
                    try:
                        page_text = page.get_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as page_error:
                        continue
                
                doc.close()
                
                if not text or len(text.strip()) < 10:
                    raise Exception("PDF file appears to be empty or contains no extractable text. The file might be image-based or corrupted.")
                
                return text
            except Exception as e:
                error_msg = str(e)
                # Continue to fallback below
        
        # Fallback: Try using pdfplumber if available
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text() or ""
                        if page_text:
                            text += page_text + "\n"
                    except Exception as page_error:
                        continue
                    
                    if not text or len(text.strip()) < 10:
                        raise Exception("PDF file appears to be empty or contains no extractable text.")
                    
                    return text
        except ImportError:
            raise Exception(
                "PDF parsing libraries not available. Please install dependencies: pip install -r requirements.txt"
            )
        except Exception as e:
            error_msg = str(e)
            error_msg_lower = error_msg.lower()  # Cache lowercased string
            # Optimized: use cached lowercased string instead of multiple .lower() calls
            if "not a pdf" in error_msg_lower or "invalid" in error_msg_lower or "cannot read" in error_msg_lower:
                raise Exception("The file is not a valid PDF document. Please upload a valid PDF file.")
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not os.path.exists(file_path):
            raise Exception(f"DOCX file not found at path: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise Exception("DOCX file is empty (0 bytes)")
        
        
        if not DOCX_AVAILABLE:
            raise Exception("python-docx is not available. Please install dependencies: pip install -r requirements.txt")
        
        try:
            doc = Document(file_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)
            
            text = "\n".join(text_parts)
            
            if not text or len(text.strip()) < 10:
                raise Exception("DOCX file appears to be empty or contains no extractable text.")
            
            return text
        except Exception as e:
            error_msg = str(e)
            error_msg_lower = error_msg.lower()  # Cache lowercased string
            # Optimized: use cached lowercased string instead of multiple .lower() calls
            if "not a docx" in error_msg_lower or "invalid" in error_msg_lower or "corrupt" in error_msg_lower or "cannot open" in error_msg_lower:
                raise Exception("The file is not a valid DOCX document. Please upload a valid DOCX file.")
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def extract_text(self, file_path: str, file_extension: str) -> str:
        """Extract text from resume file based on extension"""
        file_extension = file_extension.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def extract_skills(self, text: str) -> List[str]:
        """
        Extract skills from resume text
        Time Complexity: O(n*m) where n = text length, m = number of skills
        Space Complexity: O(k) where k = found skills (max 20)
        Optimization: Cache lowercased text, use set for O(1) membership check
        """
        text_lower = text.lower()  # Cache once
        found_skills = []
        found_skills_set = set()  # Use set for O(1) membership check
        
        for skill in self.skill_keywords:
            # Check for skill in various formats
            skill_lower = skill.lower()  # Cache skill lowercase
            # Optimized: text is already lowercased, no need for IGNORECASE flag
            pattern = rf'\b{re.escape(skill_lower)}\b'
            
            if re.search(pattern, text_lower):
                # Capitalize properly
                skill_formatted = skill.title() if '.' not in skill else skill.upper()
                if skill_formatted not in found_skills_set:
                    found_skills.append(skill_formatted)
                    found_skills_set.add(skill_formatted)
                    if len(found_skills) >= 20:  # Early exit when limit reached
                        break
        
        return found_skills[:20]  # Limit to top 20 skills
    
    def extract_experience_level(self, text: str) -> Optional[str]:
        """
        Extract experience level from resume text
        ONLY counts actual work experience, NOT projects, internships, or academic work.
        Time Complexity: O(n) where n = text length (single pass through patterns)
        Space Complexity: O(1)
        Optimization: Cache lowercased text, no IGNORECASE flag needed
        """
        text_lower = text.lower()  # Cache once
        
        # First, check for explicit fresher indicators
        fresher_pattern = r'\b(fresher|fresh\s*graduate|no\s*experience|entry\s*level|recent\s*graduate|new\s*graduate)\b'
        if re.search(fresher_pattern, text_lower):
            return "Fresher"
        
        # Look for work experience section headers
        # These keywords indicate actual work experience sections
        work_experience_section_keywords = [
            r'\b(work\s*experience|professional\s*experience|employment\s*history|work\s*history|career\s*history|experience\s*section)\b',
            r'\b(experience|employment|work\s*history)\s*:',
        ]
        
        # Check if there's a work experience section
        has_work_experience_section = False
        for pattern in work_experience_section_keywords:
            if re.search(pattern, text_lower):
                has_work_experience_section = True
                break
        
        # If no work experience section found, check for company/role patterns
        # These patterns indicate actual employment
        employment_indicators = [
            r'\b(company|employer|organization|corporation|firm|organization)\s*:',
            r'\b(worked\s*at|employed\s*at|position\s*at|role\s*at|job\s*at)\b',
            r'\b(software\s*engineer|developer|analyst|manager|engineer|consultant)\s*(?:at|in|with)\b',
        ]
        
        has_employment_indicators = False
        for pattern in employment_indicators:
            if re.search(pattern, text_lower):
                has_employment_indicators = True
                break
        
        # If no work experience section or employment indicators found, return Fresher
        if not has_work_experience_section and not has_employment_indicators:
            return "Fresher"
        
        # Now look for years of experience ONLY in work experience context
        # Patterns that explicitly mention work/professional experience
        work_experience_patterns = [
            r'\b(\d+)\s*(?:years?|yrs?|y\.?)\s*(?:of\s*)?(?:work|professional|industry|relevant)\s*experience',
            r'(?:work|professional|industry|relevant)\s*experience[:\s]+(\d+)\s*(?:years?|yrs?)',
            r'\b(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?experience\s*(?:in|with|at)\s*(?:software|development|engineering|technology)',
        ]
        
        max_years = 0
        for pattern in work_experience_patterns:
            matches = re.finditer(pattern, text_lower)
            for match in matches:
                try:
                    years = int(match.group(1))
                    max_years = max(max_years, years)
                except (IndexError, ValueError):
                    continue
        
        # If we found years in work experience context, return it
        if max_years > 0:
            return f"{max_years}yrs"
        
        # If work experience section exists but no years found, check for job titles that indicate experience
        # BUT only if we're in a work experience section context
        if has_work_experience_section or has_employment_indicators:
            # Look for senior/lead roles in work context (not project context)
            # Check if "senior" or "lead" appears near employment indicators
            senior_pattern = r'\b(senior|lead|principal|architect|manager|director)\s+(?:software|engineer|developer|analyst|consultant)'
            if re.search(senior_pattern, text_lower):
                # Verify it's in work context by checking proximity to company/employment keywords
                # This is a simplified check - if senior appears, assume 5+ years
                return "5yrs+"
        
        # Default: No valid work experience found
        return "Fresher"
    
    def extract_keywords(self, text: str) -> Dict[str, List[str]]:
        """Extract keywords including tools, technologies, and job titles"""
        text_lower = text.lower()
        keywords = {
            "tools": [],
            "technologies": [],
            "job_titles": [],
            "projects": []
        }
        
        # Extract tools and technologies (more comprehensive)
        tech_keywords = [
            # Frameworks
            'django', 'flask', 'fastapi', 'react', 'angular', 'vue', 'next.js', 'nuxt.js',
            'spring', 'express', 'laravel', 'rails', 'asp.net', '.net',
            # Tools
            'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab', 'jira', 'confluence',
            'terraform', 'ansible', 'puppet', 'chef', 'vagrant',
            # Databases
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
            'oracle', 'sqlite', 'dynamodb', 'firebase', 'supabase',
            # Cloud
            'aws', 'azure', 'gcp', 'heroku', 'vercel', 'netlify',
            # Other technologies
            'rest api', 'graphql', 'microservices', 'serverless', 'lambda',
            'machine learning', 'ai', 'data science', 'nlp', 'computer vision',
            'blockchain', 'web3', 'ethereum', 'solidity'
        ]
        
        for keyword in tech_keywords:
            # Fix: Extract replace outside f-string (backslashes not allowed in f-string expressions)
            escaped_keyword = keyword.replace(".", "\\.")
            pattern = rf'\b{re.escape(escaped_keyword)}\b'
            if re.search(pattern, text_lower, re.IGNORECASE):
                formatted = keyword.title() if '.' not in keyword else keyword
                if formatted not in keywords["technologies"]:
                    keywords["technologies"].append(formatted)
        
        # Extract job titles
        job_title_patterns = [
            r'(?:software|web|mobile|full.?stack|front.?end|back.?end|devops|data|ml|ai)\s+(?:engineer|developer|architect|specialist)',
            r'(?:senior|junior|lead|principal)\s+(?:software|web|mobile|full.?stack|front.?end|back.?end|devops|data|ml|ai)\s+(?:engineer|developer|architect)',
            r'(?:python|java|javascript|react|angular|vue|node)\s+(?:developer|engineer)',
            r'data\s+(?:scientist|engineer|analyst)',
            r'(?:machine\s+learning|ml|ai)\s+(?:engineer|scientist)',
            r'devops\s+(?:engineer|specialist)',
            r'system\s+(?:administrator|admin|architect)',
            r'qa\s+(?:engineer|tester|analyst)',
            r'product\s+(?:manager|owner)',
            r'technical\s+(?:lead|manager|architect)'
        ]
        
        for pattern in job_title_patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                title = match.group(0).title()
                if title not in keywords["job_titles"]:
                    keywords["job_titles"].append(title)
        
        # Extract project mentions (simplified - looks for "project" followed by description)
        project_patterns = [
            r'project[:\s]+([A-Z][^.!?]{10,100})',
            r'built\s+([a-z\s]{10,80})\s+(?:using|with|in)',
            r'developed\s+([a-z\s]{10,80})\s+(?:using|with|in)',
            r'created\s+([a-z\s]{10,80})\s+(?:using|with|in)'
        ]
        
        for pattern in project_patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                if len(match.groups()) > 0:
                    project = match.group(1).strip()[:80]
                    if project and project not in keywords["projects"]:
                        keywords["projects"].append(project)
        
        return keywords
    
    def extract_name(self, text: str) -> Optional[str]:
        """
        Extract name from resume text (usually at the beginning)
        Time Complexity: O(n) where n = number of lines checked (max 10)
        Space Complexity: O(1)
        Optimization: Cache excluded words set, optimize regex checks
        """
        lines = text.split('\n')[:10]  # Check first 10 lines
        # Optimized: cache excluded words in a set for O(1) lookup
        excluded_words = {'email', 'phone', 'address', 'resume', 'cv'}
        
        for line in lines:
            line = line.strip()
            if len(line) > 3 and len(line) < 50:  # Reasonable name length
                # Check if it looks like a name (contains letters, may have spaces)
                if re.match(r'^[A-Za-z\s\.\-]+$', line):
                    # Optimized: use set for O(1) membership check instead of list
                    line_words = {word.lower() for word in line.split()}
                    if not line_words.intersection(excluded_words):
                        # Check if it's not an email or phone
                        if '@' not in line and not re.match(r'^[\d\s\-\+\(\)]+$', line):
                            return line.title()
        return None
    
    def extract_email(self, text: str) -> Optional[str]:
        """Extract email address from resume text"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)
        if matches:
            return matches[0].lower()
        return None
    
    def extract_projects(self, text: str) -> List[Dict[str, str]]:
        """
        Extract project information from resume text.
        Extracts from both PROJECTS sections and Internship sections.
        Preserves bullet formatting, excludes internship metadata, and properly formats output.
        """
        projects = []
        text_lines = text.split('\n')
        text_lower = text.lower()
        
        # Major section headers that indicate end of current section
        section_end_keywords = [
            'training and certification', 'certification', 'certifications', 'training',
            'technical skills', 'skills', 'technical expertise',
            'education', 'academic background', 'qualifications',
            'work experience', 'employment', 'professional experience',
            'achievements', 'awards', 'honors',
            'languages', 'language proficiency',
            'references', 'contact', 'personal information',
            'summary', 'objective', 'profile'
        ]
        
        def clean_project_description(description_lines: List[str], is_internship: bool) -> Dict[str, Any]:
            """
            Clean and format project description, preserving bullet structure.
            Returns structured data with tech_stack, responsibilities, and tools as separate arrays.
            Dynamically detects content types without hard-coding keywords.
            """
            tech_stack_items = []
            tools_items = []
            responsibility_bullets = []
            other_lines = []
            
            # Patterns to identify and exclude
            internship_metadata_patterns = [
                r'^(frontend|backend|full.?stack|software|web|mobile)\s+(developer|engineer|intern)',
                r'\([^)]*\d{4}[^)]*\)',  # Dates in parentheses
                r'^(january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{4}',
                r'organization[:\s]+',
                r'company[:\s]+',
                r'employer[:\s]+',
                r'^[A-Z\s]{2,30}$',  # All caps section headers
            ]
            
            # Technology/framework patterns for dynamic detection
            tech_term_patterns = [
                r'\b(react|angular|vue|node|javascript|typescript|python|java|html|css|django|flask|fastapi|express|spring|laravel|rails|\.net|\.js|\.py|\.ts|\.jsx|\.tsx|api|rest|graphql|sql|mongodb|postgresql|mysql|redis|docker|kubernetes|aws|azure|gcp)\b',
            ]
            
            # Tool patterns for dynamic detection
            tool_term_patterns = [
                r'\b(git|github|gitlab|vscode|visual\s+studio|postman|jira|confluence|jenkins|terraform|ansible|puppet|chef|vagrant)\b',
            ]
            
            # Bullet markers (comprehensive list including unicode)
            bullet_markers = ['-', '•', '*', '·', '▪', '▸', '▹', '▪', '▫', '◦', '‣', '⁃', '⁌', '⁍', '→', '➜', '➤', '○', '●']
            
            # Action verbs that indicate responsibilities
            action_verbs = ['developed', 'created', 'built', 'designed', 'implemented', 'worked', 
                          'collaborated', 'enhanced', 'practiced', 'gained', 'integrated', 'used',
                          'managed', 'led', 'improved', 'optimized', 'delivered', 'established',
                          'configured', 'deployed', 'maintained', 'debugged', 'tested', 'wrote']
            
            for line in description_lines:
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                
                line_lower = line_stripped.lower()
                
                # Skip section headers
                is_section_header = False
                for keyword in section_end_keywords:
                    if re.match(rf'^{re.escape(keyword)}[:\s]*$', line_lower) or re.match(rf'^{re.escape(keyword.upper())}[:\s]*$', line_stripped):
                        is_section_header = True
                        break
                if is_section_header:
                    continue
                
                # Skip internship metadata
                if is_internship:
                    is_metadata = False
                    for pattern in internship_metadata_patterns:
                        if re.search(pattern, line_stripped, re.IGNORECASE):
                            is_metadata = True
                            break
                    if is_metadata:
                        continue
                
                # Detect tech stack dynamically using semantic analysis
                is_tech_stack = False
                if ':' in line_stripped:
                    parts = line_stripped.split(':', 1)
                    if len(parts) == 2:
                        prefix = parts[0].lower().strip()
                        content = parts[1].strip()
                        
                        # Check if prefix suggests tech/technologies (dynamic pattern matching)
                        tech_prefix_indicators = ['tech', 'technology', 'technologies', 'stack', 'framework', 'library', 'libraries', 'language', 'languages']
                        has_tech_prefix = any(indicator in prefix for indicator in tech_prefix_indicators)
                        
                        # Check if content contains technology terms
                        has_tech_terms = any(re.search(pattern, content, re.IGNORECASE) for pattern in tech_term_patterns)
                        
                        # Check if content looks like a list (comma-separated or pipe-separated)
                        is_list_like = ',' in content or '|' in content or len(content.split()) <= 10
                        
                        if has_tech_prefix and has_tech_terms and is_list_like and len(content) > 5:
                            # Extract tech items
                            tech_items = [item.strip() for item in re.split(r'[,|]', content) if item.strip()]
                            if tech_items:
                                tech_stack_items.extend(tech_items)
                                is_tech_stack = True
                
                # Detect tools dynamically using semantic analysis
                is_tools = False
                if ':' in line_stripped:
                    parts = line_stripped.split(':', 1)
                    if len(parts) == 2:
                        prefix = parts[0].lower().strip()
                        content = parts[1].strip()
                        
                        # Check if prefix suggests tools (dynamic pattern matching)
                        tools_prefix_indicators = ['tool', 'tools', 'software', 'platform', 'environment', 'ide', 'editor']
                        has_tools_prefix = any(indicator in prefix for indicator in tools_prefix_indicators)
                        
                        # Check if content contains tool terms
                        has_tool_terms = any(re.search(pattern, content, re.IGNORECASE) for pattern in tool_term_patterns)
                        
                        # Check if content looks like a list
                        is_list_like = ',' in content or '|' in content or len(content.split()) <= 10
                        
                        if has_tools_prefix and has_tool_terms and is_list_like and len(content) > 3:
                            # Extract tool items
                            tool_items = [item.strip() for item in re.split(r'[,|]', content) if item.strip()]
                            if tool_items:
                                tools_items.extend(tool_items)
                                is_tools = True
                
                if is_tech_stack or is_tools:
                    continue
                
                # Detect and preserve bullet points
                is_bullet = False
                for marker in bullet_markers:
                    # Check if line starts with bullet marker (with optional whitespace)
                    if re.match(rf'^\s*[{re.escape(marker)}]', line_stripped):
                        is_bullet = True
                        # Extract content after bullet marker
                        content_after_bullet = re.sub(rf'^\s*[{re.escape(marker)}]\s*', '', line_stripped).strip()
                        if content_after_bullet:
                            # Store bullet content without marker (we'll format it later)
                            responsibility_bullets.append(content_after_bullet)
                        break
                
                if not is_bullet:
                    # Check if line might be a bullet without marker (starts with action verb)
                    # This handles cases where PDF extraction loses bullet markers
                    words = line_stripped.split()
                    if words and words[0].lower() in action_verbs and len(line_stripped) > 20:
                        # Likely a responsibility bullet without marker
                        responsibility_bullets.append(line_stripped)
                    elif len(line_stripped) > 10 and not line_stripped.endswith(('.', ':', ';')):
                        # Check if it's a short descriptive line that might be a bullet
                        if len(line_stripped.split()) <= 15:
                            responsibility_bullets.append(line_stripped)
                        else:
                            other_lines.append(line_stripped)
                    else:
                        # Other descriptive line
                        other_lines.append(line_stripped)
            
            # Build formatted summary string with HTML line breaks for proper display
            formatted_parts = []
            
            # Add tech stack section if available
            if tech_stack_items:
                tech_content = ', '.join(tech_stack_items)
                formatted_parts.append(f"Tech Stack: {tech_content}")
            
            # Add responsibilities (bullets) - each on its own line with HTML breaks
            if responsibility_bullets:
                # Format each bullet with proper spacing
                for bullet in responsibility_bullets:
                    formatted_parts.append(f"• {bullet}")
            
            # Add other descriptive lines
            if other_lines:
                formatted_parts.extend(other_lines)
            
            # Add tools section at end if available
            if tools_items:
                tools_content = ', '.join(tools_items)
                formatted_parts.append(f"Tools: {tools_content}")
            
            # Join with HTML line breaks for proper rendering in frontend
            formatted_summary = '<br>'.join(formatted_parts)
            
            # Return structured data
            return {
                "summary": formatted_summary,  # HTML-formatted string for display
                "tech_stack": tech_stack_items,  # Array of tech items
                "responsibilities": responsibility_bullets,  # Array of bullet point texts
                "tools": tools_items,  # Array of tool items
                "raw_text": '\n'.join(description_lines)  # Original text for fallback
            }
        
        def extract_from_section(section_start_idx: int, section_name: str) -> List[Dict[str, str]]:
            """Extract projects from a specific section"""
            section_projects = []
            section_lines = []
            
            # Extract content from section until next major section
            for i in range(section_start_idx + 1, len(text_lines)):
                line = text_lines[i].strip()
                if not line:
                    section_lines.append('')  # Preserve empty lines for project separation
                    continue
                
                line_lower = line.lower()
                
                # Check if we've hit a new major section - stop extraction
                is_section_end = False
                for keyword in section_end_keywords:
                    if (re.match(rf'^{re.escape(keyword)}[:\s]*$', line_lower) or
                        re.match(rf'^{re.escape(keyword.upper())}[:\s]*$', line)):
                        is_section_end = True
                        break
                
                if is_section_end:
                    break
                
                section_lines.append(text_lines[i])
            
            # Parse projects from section lines
            current_project = None
            current_description_lines = []
            is_internship_section = 'internship' in section_name.lower()
            
            # Patterns to identify internship role titles (to skip them)
            role_title_patterns = [
                r'^(frontend|backend|full.?stack|software|web|mobile|data|devops|qa|test)\s+(developer|engineer|intern|specialist|analyst|architect)',
                r'\s+(developer|engineer|intern|specialist|analyst|architect)\s*[-–]',  # Role before dash
            ]
            
            # Patterns to identify dates and organization lines (to skip)
            date_patterns = [
                r'\([^)]*\d{4}[^)]*\)',  # Dates in parentheses
                r'^(january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{4}',
                r'organization[:\s]+',
                r'company[:\s]+',
                r'employer[:\s]+',
            ]
            
            for idx, line in enumerate(section_lines):
                line_stripped = line.strip()
                
                if not line_stripped:
                    # Empty line - if we have a project, save it
                    if current_project and current_description_lines:
                        cleaned_data = clean_project_description(current_description_lines, is_internship_section)
                        if cleaned_data.get("summary"):  # Only add if has content
                            # Merge structured data into project
                            current_project.update(cleaned_data)
                            section_projects.append(current_project)
                        current_project = None
                        current_description_lines = []
                    continue
                
                line_lower = line_stripped.lower()
                is_bullet = line_stripped.startswith(('-', '•', '*', '·'))
                line_no_bullet = re.sub(r'^[\s]*[-•*·]\s*', '', line_stripped).strip()
                
                # Skip internship role titles and metadata in internship sections
                if is_internship_section:
                    # Check if this is a role title line
                    is_role_title = False
                    for pattern in role_title_patterns:
                        if re.search(pattern, line_stripped, re.IGNORECASE):
                            is_role_title = True
                            break
                    
                    # Check if this is a date or organization line
                    is_metadata_line = False
                    for pattern in date_patterns:
                        if re.search(pattern, line_stripped, re.IGNORECASE):
                            is_metadata_line = True
                            break
                    
                    # Skip role titles and metadata lines - don't use them as project titles
                    if is_role_title or is_metadata_line:
                        continue
                
                # Check if line is a project title
                is_project_title = False
                
                if not is_bullet:
                    words = line_no_bullet.split()
                    word_count = len(words)
                    
                    # Project title indicators (dynamic - no hard-coding)
                    title_indicators = ['app', 'application', 'system', 'platform', 'project', 'web', 'mobile',
                                       'quiz', 'ordering', 'food', 'pharma', 'game', 'dashboard', 'portal',
                                       'website', 'site', 'tool', 'service', 'api']
                    
                    # Action verbs that indicate description, not title
                    action_verbs = ['developed', 'created', 'built', 'designed', 'implemented', 
                                   'worked', 'collaborated', 'enhanced', 'practiced', 'gained',
                                   'integrated', 'used', 'tools', 'tech', 'demonstrated', 'managed']
                    
                    # Role-related words that indicate this is NOT a project title
                    role_words = ['intern', 'developer', 'engineer', 'specialist', 'analyst', 'architect', 'manager']
                    
                    # Check for project title patterns
                    if (word_count >= 2 and word_count <= 10 and
                        re.match(r'^[A-Z]', line_no_bullet) and
                        not any(word.lower() in action_verbs for word in words[:3]) and
                        not any(word.lower() in role_words for word in words)):  # Exclude role words
                        
                        # Check for title indicators
                        if any(indicator in line_no_bullet.lower() for indicator in title_indicators):
                            is_project_title = True
                        # Or short capitalized phrase without punctuation
                        elif (word_count <= 6 and 
                              not line_no_bullet.endswith(('.', '!', '?')) and
                              ':' not in line_no_bullet):
                            is_project_title = True
                
                if is_project_title:
                    # Save previous project if exists
                    if current_project and current_description_lines:
                        cleaned_data = clean_project_description(current_description_lines, is_internship_section)
                        if cleaned_data.get("summary"):
                            # Merge structured data into project
                            current_project.update(cleaned_data)
                            section_projects.append(current_project)
                    
                    # Start new project
                    current_project = {
                        "name": line_no_bullet[:100],
                        "summary": "",
                        "tech_stack": [],
                        "responsibilities": [],
                        "tools": []
                    }
                    current_description_lines = []
                elif current_project:
                    # Description line - preserve formatting
                    current_description_lines.append(line_stripped)
                elif is_internship_section and not current_project:
                    # In internship section, look for project mentions in bullet points
                    # Extract full project name from bullet points like "Developed a Pharma Quiz Web Application using..."
                    # Pattern: "Developed a [Full Project Name] using/with/for"
                    # Use a more flexible pattern that captures everything until "using/with/for"
                    pattern = r'(?:developed|created|built|designed)\s+(?:a\s+)?([A-Z][^.!?]*?)\s+(?:using|with|for)\s+'
                    match = re.search(pattern, line_no_bullet, re.IGNORECASE)
                    
                    if match:
                        potential_project = match.group(1).strip()
                        # Clean up: remove trailing words that aren't part of project name
                        # Keep only up to project name endings
                        project_endings = ['web application', 'web app', 'application', 'app', 'system', 'platform', 'project']
                        found_ending = False
                        for ending in project_endings:
                            ending_lower = ending.lower()
                            if ending_lower in potential_project.lower():
                                # Find the ending and keep everything up to and including it
                                ending_pos = potential_project.lower().find(ending_lower)
                                if ending_pos >= 0:
                                    # Extract from start to end of the ending phrase
                                    words_before = potential_project[:ending_pos].strip().split()
                                    # Reconstruct with proper capitalization
                                    if words_before:
                                        potential_project = ' '.join(words_before) + ' ' + ending.title()
                                    else:
                                        potential_project = ending.title()
                                    found_ending = True
                                    break
                        
                        # If no ending found, try to extract a reasonable project name
                        if not found_ending:
                            # Take first few capitalized words (likely the project name)
                            words = potential_project.split()
                            if len(words) >= 2:
                                # Take up to 6 words as project name
                                potential_project = ' '.join(words[:6])
                        
                        # Verify it's a real project name (not a role or action)
                        title_indicators = ['app', 'application', 'quiz', 'ordering', 'food', 'pharma', 'web', 'system', 'platform']
                        role_words = ['intern', 'developer', 'engineer', 'specialist']
                        
                        # Check if it contains project indicators and doesn't contain role words
                        has_project_indicator = any(indicator in potential_project.lower() for indicator in title_indicators)
                        has_role_word = any(role in potential_project.lower() for role in role_words)
                        
                        if has_project_indicator and not has_role_word and len(potential_project.split()) >= 2:
                            # Start new project with full name
                            current_project = {
                                "name": potential_project[:100],
                                "summary": "",
                                "tech_stack": [],
                                "responsibilities": [],
                                "tools": []
                            }
                            current_description_lines = [line_stripped]
            
            # Save the last project
            if current_project and current_description_lines:
                cleaned_data = clean_project_description(current_description_lines, is_internship_section)
                if cleaned_data.get("summary"):
                    # Merge structured data into project
                    current_project.update(cleaned_data)
                    section_projects.append(current_project)
            
            return section_projects
        
        # Extract from PROJECTS section
        projects_section_start = None
        for i, line in enumerate(text_lines):
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            if re.match(r'^(projects?|project\s+experience|personal\s+projects?|notable\s+projects?)[:\s]*$', line_lower):
                projects_section_start = i
                break
        
        if projects_section_start is not None:
            projects.extend(extract_from_section(projects_section_start, "PROJECTS"))
        
        # Extract from Internship sections
        for i, line in enumerate(text_lines):
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            if re.match(r'^(internship\s+experience|internships?)[:\s]*$', line_lower):
                # Extract projects from this internship section
                internship_projects = extract_from_section(i, "INTERNSHIP")
                projects.extend(internship_projects)
                # Only process first internship section to avoid duplicates
                break
        
        # Deduplicate projects based on name similarity
        unique_projects = []
        seen_names = set()
        
        for proj in projects:
            if not proj.get("name") or not proj.get("summary"):
                continue
            
            name = proj["name"].strip()
            name_lower = name.lower()
            
            # Check for duplicates
            is_duplicate = False
            for seen_name in seen_names:
                seen_lower = seen_name.lower()
                # Exact match
                if name_lower == seen_lower:
                    is_duplicate = True
                    break
                # Fuzzy matching: check if names are very similar
                if name_lower in seen_lower or seen_lower in name_lower:
                    shorter = min(len(name_lower), len(seen_lower))
                    longer = max(len(name_lower), len(seen_lower))
                    if shorter > 0 and shorter / longer > 0.7:
                        is_duplicate = True
                        break
                # Check word overlap
                name_words = set(name_lower.split())
                seen_words = set(seen_lower.split())
                if len(name_words) > 0 and len(seen_words) > 0:
                    overlap = len(name_words & seen_words) / max(len(name_words), len(seen_words))
                    if overlap > 0.6:
                        is_duplicate = True
                        break
            
            if not is_duplicate:
                seen_names.add(name)
                unique_projects.append(proj)
        
        return unique_projects
    
    def categorize_skills(self, skills: List[str], text: str) -> Dict[str, List[str]]:
        """Categorize skills into Programming, AI/ML, Tools, Soft Skills"""
        text_lower = text.lower()
        categorized = {
            "Programming": [],
            "AI/ML": [],
            "Tools": [],
            "Soft Skills": []
        }
        
        # Programming languages and frameworks
        programming_keywords = [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'php', 'ruby',
            'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'html', 'css', 'sass', 'less',
            'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'fastapi',
            'spring', 'laravel', 'rails', 'asp.net', '.net', 'next.js', 'nuxt.js'
        ]
        
        # AI/ML keywords
        ai_ml_keywords = [
            'machine learning', 'ml', 'ai', 'artificial intelligence', 'deep learning',
            'neural network', 'tensorflow', 'pytorch', 'keras', 'scikit-learn',
            'nlp', 'natural language processing', 'computer vision', 'cv',
            'data science', 'data analysis', 'pandas', 'numpy', 'opencv'
        ]
        
        # Tools and DevOps
        tools_keywords = [
            'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab', 'jira',
            'terraform', 'ansible', 'linux', 'bash', 'shell', 'ci/cd',
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch',
            'aws', 'azure', 'gcp', 'heroku', 'vercel', 'netlify'
        ]
        
        # Soft skills (extract from text, not from skill list typically)
        soft_skills_keywords = [
            'leadership', 'communication', 'teamwork', 'collaboration', 'problem solving',
            'agile', 'scrum', 'project management', 'mentoring', 'presentation'
        ]
        
        # Categorize each skill
        for skill in skills:
            skill_lower = skill.lower()
            categorized_flag = False
            
            # Check programming
            for keyword in programming_keywords:
                if keyword in skill_lower or skill_lower in keyword:
                    categorized["Programming"].append(skill)
                    categorized_flag = True
                    break
            
            if not categorized_flag:
                # Check AI/ML
                for keyword in ai_ml_keywords:
                    if keyword in skill_lower or skill_lower in keyword:
                        categorized["AI/ML"].append(skill)
                        categorized_flag = True
                        break
            
            if not categorized_flag:
                # Check tools
                for keyword in tools_keywords:
                    if keyword in skill_lower or skill_lower in keyword:
                        categorized["Tools"].append(skill)
                        categorized_flag = True
                        break
            
            # If not categorized, add to Programming as default
            if not categorized_flag:
                categorized["Programming"].append(skill)
        
        # Extract additional soft skills from text
        for keyword in soft_skills_keywords:
            if keyword in text_lower:
                formatted = keyword.title()
                if formatted not in categorized["Soft Skills"]:
                    categorized["Soft Skills"].append(formatted)
        
        # Remove duplicates within each category
        for category in categorized:
            categorized[category] = list(dict.fromkeys(categorized[category]))
        
        return categorized
    
    def generate_interview_topics(self, skills: List[str], keywords: Dict[str, List[str]], 
                                   projects: List[Dict[str, str]], text: str) -> List[str]:
        """Generate AI interview preparation topics based on resume content"""
        topics = set()
        text_lower = text.lower()
        
        # Core topics based on skills
        skill_to_topic = {
            'python': 'Python',
            'java': 'Java',
            'javascript': 'JavaScript',
            'typescript': 'TypeScript',
            'react': 'React',
            'angular': 'Angular',
            'vue': 'Vue.js',
            'node.js': 'Node.js',
            'django': 'Django',
            'flask': 'Flask',
            'fastapi': 'FastAPI',
            'sql': 'SQL',
            'postgresql': 'PostgreSQL',
            'mysql': 'MySQL',
            'mongodb': 'MongoDB',
            'redis': 'Redis',
            'docker': 'Docker',
            'kubernetes': 'Kubernetes',
            'aws': 'AWS',
            'azure': 'Azure',
            'gcp': 'GCP',
            'git': 'Git',
            'machine learning': 'Machine Learning',
            'ml': 'Machine Learning',
            'ai': 'Artificial Intelligence',
            'nlp': 'NLP',
            'data science': 'Data Science',
            'computer vision': 'Computer Vision',
            'tensorflow': 'TensorFlow',
            'pytorch': 'PyTorch'
        }
        
        # Add topics from skills
        for skill in skills:
            skill_lower = skill.lower()
            for key, topic in skill_to_topic.items():
                if key in skill_lower:
                    topics.add(topic)
        
        # Add topics from technologies
        for tech in keywords.get("technologies", []):
            tech_lower = tech.lower()
            for key, topic in skill_to_topic.items():
                if key in tech_lower:
                    topics.add(topic)
        
        # Add common interview topics based on domain
        if any(term in text_lower for term in ['machine learning', 'ml', 'ai', 'data science', 'nlp']):
            topics.update(['Machine Learning', 'Statistics', 'Linear Algebra', 'Probability'])
            if 'nlp' in text_lower or 'natural language' in text_lower:
                topics.add('NLP')
            if 'computer vision' in text_lower or 'cv' in text_lower:
                topics.add('Computer Vision')
        
        if any(term in text_lower for term in ['web', 'frontend', 'backend', 'full stack']):
            topics.update(['System Design', 'REST APIs', 'HTTP/HTTPS', 'Web Security'])
        
        if any(term in text_lower for term in ['database', 'sql', 'postgresql', 'mysql', 'mongodb']):
            topics.update(['Database Design', 'SQL Queries', 'Indexing', 'Transactions'])
        
        if any(term in text_lower for term in ['cloud', 'aws', 'azure', 'gcp']):
            topics.add('Cloud Architecture')
        
        if any(term in text_lower for term in ['docker', 'kubernetes', 'devops']):
            topics.update(['DevOps', 'CI/CD', 'Containerization'])
        
        # Always add fundamental topics
        topics.update(['Data Structures', 'Algorithms', 'Problem Solving'])
        
        # Add Git if mentioned
        if 'git' in text_lower:
            topics.add('Git')
        
        # Add APIs if mentioned
        if any(term in text_lower for term in ['api', 'rest', 'graphql']):
            topics.add('APIs')
        
        return sorted(list(topics))[:20]  # Limit to top 20 topics
    
    def calculate_resume_rating(self, name: Optional[str], email: Optional[str], 
                                skills: List[str], projects: List[Dict[str, str]],
                                experience_level: Optional[str], text_length: int,
                                keywords: Dict[str, List[str]]) -> float:
        """Calculate resume rating out of 5 based on multiple factors"""
        score = 0.0
        max_score = 5.0
        
        # Clarity (0.5 points) - name and email present
        if name:
            score += 0.25
        if email:
            score += 0.25
        
        # Skills (1.5 points) - based on number and diversity
        skills_score = min(len(skills) / 15.0, 1.0) * 1.5  # Max 1.5 for 15+ skills
        score += skills_score
        
        # Projects (1.0 point) - based on number and detail
        projects_score = min(len(projects) / 5.0, 1.0) * 1.0  # Max 1.0 for 5+ projects
        score += projects_score
        
        # Formatting/Completeness (1.0 point) - based on text length and structure
        if text_length > 1000:
            score += 0.5
        if text_length > 2000:
            score += 0.3
        if len(keywords.get("job_titles", [])) > 0:
            score += 0.2
        
        # Experience level (1.0 point) - if specified
        if experience_level and experience_level not in ["Not specified", "Unknown", "Fresher"]:
            score += 0.5
        elif experience_level == "Fresher":
            score += 0.2
        
        # Ensure score is between 0 and 5
        return round(min(max(score, 0.0), max_score), 2)
    
    def generate_resume_summary(self, name: Optional[str], email: Optional[str],
                               skills: List[str], experience_level: Optional[str],
                               keywords: Dict[str, List[str]], text: str) -> str:
        """Generate a clear resume summary paragraph based ONLY on actual resume content"""
        summary_parts = []
        
        # Start with name if available
        if name:
            summary_parts.append(f"{name} is")
        else:
            summary_parts.append("This candidate is")
        
        # Add experience level
        if experience_level and experience_level not in ["Not specified", "Unknown"]:
            if experience_level == "Fresher":
                summary_parts.append("a fresher")
            else:
                summary_parts.append(f"an experienced professional with {experience_level} of experience")
        else:
            summary_parts.append("a professional")
        
        # Add domain/role ONLY from extracted job titles (actual resume content)
        job_titles = keywords.get("job_titles", [])
        if job_titles:
            primary_role = job_titles[0]
            summary_parts.append(f"specializing in {primary_role.lower()}")
        
        # Build expertise description from ACTUAL extracted skills and technologies only
        # NO hard-coded assumptions or generic fallbacks
        expertise_parts = []
        
        # Use actual extracted technologies (not assumptions)
        technologies = keywords.get("technologies", [])
        if technologies:
            # Group technologies by domain based on what's actually present
            tech_list = technologies[:5]  # Limit to top 5
            tech_str = ", ".join(tech_list[:3])
            if tech_str:
                expertise_parts.append(f"proficient in {tech_str}")
        
        # Use actual extracted skills (not assumptions)
        if skills:
            # Get primary skills (limit to top 3-5 most relevant)
            primary_skills = skills[:5]
            skill_str = ", ".join(primary_skills[:3])
            if skill_str and not technologies:  # Only add if we don't have technologies
                expertise_parts.append(f"skilled in {skill_str}")
        
        # Add expertise description only if we have actual data
        if expertise_parts:
            summary_parts.append(f"with {', '.join(expertise_parts)}")
        
        # Add key strengths based on actual resume content
        strengths = []
        if len(skills) >= 10:
            strengths.append("strong technical skills")
        if len(keywords.get("technologies", [])) >= 5:
            strengths.append("diverse technology experience")
        if keywords.get("job_titles"):
            strengths.append("relevant industry experience")
        
        # Add projects if available (actual resume content)
        projects = keywords.get("projects", [])
        if projects and len(projects) > 0:
            strengths.append("project experience")
        
        if strengths:
            summary_parts.append(f"demonstrating {', '.join(strengths)}")
        
        summary = " ".join(summary_parts) + "."
        return summary
    
    def generate_enhanced_summary(self, parsed_data: Dict[str, Any], text: str) -> Dict[str, Any]:
        """Generate comprehensive enhanced summary with all required components"""
        name = parsed_data.get("name")
        email = parsed_data.get("email")
        skills = parsed_data.get("skills", [])
        experience_level = parsed_data.get("experience_level")
        keywords = parsed_data.get("keywords", {})
        text_length = parsed_data.get("text_length", 0)
        
        # Extract projects
        projects = self.extract_projects(text)
        
        # Categorize skills
        categorized_skills = self.categorize_skills(skills, text)
        
        # Generate resume summary paragraph
        resume_summary = self.generate_resume_summary(name, email, skills, experience_level, keywords, text)
        
        # Generate interview topics
        interview_topics = self.generate_interview_topics(skills, keywords, projects, text)
        
        # Calculate rating
        rating = self.calculate_resume_rating(name, email, skills, projects, experience_level, text_length, keywords)
        
        return {
            "resume_summary": resume_summary,
            "skills_summary": categorized_skills,
            "projects_summary": projects,
            "interview_topics": interview_topics,
            "resume_rating": rating
        }
    
    def generate_interview_modules(self, parsed_data: Dict[str, Any], text: str) -> Dict[str, Any]:
        """Generate Interview Modules Overview with 4 modules - skills split based on resume"""
        skills = parsed_data.get("skills", [])
        experience_level = parsed_data.get("experience_level", "Fresher")
        keywords = parsed_data.get("keywords", {})
        # Get projects from summary if available, otherwise extract them
        summary = parsed_data.get("summary", {})
        projects = summary.get("projects_summary", []) if summary else []
        if not projects:
            projects = self.extract_projects(text)
        text_lower = text.lower()
        
        # Split skills into modules based on resume content
        technical_skills_result = self._extract_technical_skills(skills, keywords, text_lower, text, projects)
        core_technical_skills = technical_skills_result.get("core_technical_skills", [])
        domain_skills = technical_skills_result.get("domain_skills", [])
        
        # Combine core technical skills and domain skills for display
        all_technical_skills = core_technical_skills + domain_skills
        total_skills_count = len(all_technical_skills)
        
        coding_topics = self._extract_coding_topics(skills, keywords, text_lower)
        hr_skills = self._extract_hr_skills(text_lower, keywords)
        star_points = self._extract_star_points(text, projects, keywords)
        
        # Technical Interview Module - core technical & domain skills
        technical_description = f"Focus on {total_skills_count} core technical skills and domain expertise extracted from your resume."
        
        # Coding / Online Test Module
        difficulty_level = self._determine_coding_difficulty(experience_level, skills, projects, text_lower)
        coding_platforms = self._recommend_coding_platforms(skills, text_lower)
        
        # HR Interview Module
        hr_evaluation_points = self._generate_hr_evaluation_points()
        
        # Behavioral Interview Module
        star_guidance = self._generate_star_guidance()
        
        return {
            "technical_interview": {
                "description": technical_description,
                "skills": all_technical_skills  # Combined core + domain skills
            },
            "coding_test": {
                "difficulty_level": difficulty_level,
                "platforms": coding_platforms,
                "topics": coding_topics
            },
            "hr_interview": {
                "evaluation_points": hr_evaluation_points,
                "skills": hr_skills
            },
            "behavioral_interview": {
                "star_guidance": star_guidance,
                "star_points": star_points
            }
        }
    
    def _extract_technical_skills(self, skills: List[str], keywords: Dict[str, List[str]], 
                                  text_lower: str, text: str, projects: List[Dict[str, str]]) -> Dict[str, List[str]]:
        """
        Extract core technical and domain skills for Technical Interview.
        Dynamically extracts ONLY from resume content - no hard-coded skill lists.
        Returns: {
            "core_technical_skills": List[str],  # Languages, frameworks, tools, platforms
            "domain_skills": List[str]  # Domain expertise like "Web Development", "API Integration"
        }
        """
        # Use dict to track normalized -> best form mapping for deduplication
        normalized_to_best_form = {}  # normalized_key -> best_form
        domain_skills_set = set()
        
        # Normalization function for deduplication
        def normalize_skill(term: str) -> str:
            """
            Normalize a skill term for deduplication.
            Handles variations like: Node.js/Node/NODE.JS -> nodejs
            """
            if not term:
                return ""
            
            # Convert to lowercase
            normalized = term.lower().strip()
            
            # Remove common punctuation and special chars
            normalized = re.sub(r'[^\w\s]', '', normalized)
            
            # Remove extra whitespace
            normalized = re.sub(r'\s+', '', normalized)
            
            # Handle common variations
            # GitHub/GitHub/Git -> git
            if normalized.startswith('github'):
                normalized = 'github'
            elif normalized == 'git':
                normalized = 'git'
            
            # Node.js/Node/NODE.JS -> nodejs
            if normalized.startswith('node'):
                normalized = 'nodejs'
            
            # React.js/React/ReactJS -> react
            if normalized.startswith('react'):
                normalized = 'react'
            
            # JavaScript/JS/Javascript -> javascript
            if normalized in ['js', 'javascript']:
                normalized = 'javascript'
            
            # HTML -> html
            if normalized == 'html':
                normalized = 'html'
            
            # CSS -> css
            if normalized == 'css':
                normalized = 'css'
            
            # TypeScript/TS/Typescript -> typescript
            if normalized in ['ts', 'typescript']:
                normalized = 'typescript'
            
            # Redux -> redux
            if normalized == 'redux':
                normalized = 'redux'
            
            # Tailwind CSS/Tailwind -> tailwindcss
            if 'tailwind' in normalized:
                normalized = 'tailwindcss'
            
            # Bootstrap -> bootstrap
            if normalized == 'bootstrap':
                normalized = 'bootstrap'
            
            # React Query/TanStack Query/TanStack -> reactquery
            if 'reactquery' in normalized or 'tanstackquery' in normalized or 'tanstack' in normalized:
                normalized = 'reactquery'
            elif 'react' in normalized and 'query' in normalized:
                normalized = 'reactquery'
            
            return normalized
        
        def get_best_form(term: str) -> str:
            """
            Get the best formatted version of a skill term.
            Prefers: Proper capitalization, standard naming conventions
            """
            if not term:
                return ""
            
            term_clean = term.strip()
            term_lower = term_clean.lower()
            
            # Handle known technologies with preferred formatting
            preferred_forms = {
                'git': 'Git',
                'github': 'GitHub',
                'gitlab': 'GitLab',
                'nodejs': 'Node.js',
                'node': 'Node.js',
                'react': 'React',
                'reactjs': 'React',
                'react.js': 'React',
                'javascript': 'JavaScript',
                'js': 'JavaScript',
                'typescript': 'TypeScript',
                'ts': 'TypeScript',
                'html': 'HTML',
                'css': 'CSS',
                'redux': 'Redux',
                'tailwindcss': 'Tailwind CSS',
                'tailwind': 'Tailwind CSS',
                'bootstrap': 'Bootstrap',
                'reactquery': 'React Query',
                'tanstack': 'TanStack Query',
                'api': 'API',
                'rest': 'REST API',
                'graphql': 'GraphQL',
            }
            
            normalized = normalize_skill(term_clean)
            if normalized in preferred_forms:
                return preferred_forms[normalized]
            
            # If it's an acronym (all caps, 2-5 chars), keep it uppercase
            if re.match(r'^[A-Z]{2,5}$', term_clean):
                return term_clean
            
            # If it contains a dot (like React.js), preserve the format
            if '.' in term_clean and term_clean[0].isupper():
                return term_clean
            
            # Title case for multi-word terms
            if ' ' in term_clean or '-' in term_clean:
                # Handle special cases
                if 'css' in term_lower:
                    return term_clean.replace('css', 'CSS').replace('Css', 'CSS')
                if 'api' in term_lower:
                    return term_clean.replace('api', 'API').replace('Api', 'API')
                return term_clean.title()
            
            # Capitalize first letter for single words
            return term_clean.capitalize()
        
        # Soft skills patterns to exclude
        soft_skill_patterns = [
            r'\b(communication|teamwork|leadership|collaboration|problem solving|problem-solving|'
            r'adaptability|time management|presentation|interpersonal|negotiation|mentoring|'
            r'project management|agile|scrum|kanban|critical thinking|analytical thinking|'
            r'creativity|innovation|work ethic|professionalism|multitasking|organization|'
            r'attention to detail|detail-oriented|self-motivated|proactive|flexible|'
            r'customer service|client relations|stakeholder management|conflict resolution)\b'
        ]
        
        def is_soft_skill(term: str) -> bool:
            """Check if a term is a soft skill"""
            term_lower = term.lower()
            for pattern in soft_skill_patterns:
                if re.search(pattern, term_lower):
                    return True
            return False
        
        # Invalid standalone terms (not skills by themselves)
        invalid_standalone_terms = {
            'component', 'components', 'integration', 'integrations', 'props', 'prop',
            'hook', 'hooks', 'state', 'context', 'reusability', 'responsive',
            'development', 'design', 'framework', 'library', 'libraries', 'tool', 'tools',
            'platform', 'platforms', 'system', 'systems', 'service', 'services',
            'using', 'with', 'built', 'developed', 'created', 'implemented', 'used',
            'and', 'or', 'the', 'a', 'an', 'in', 'on', 'at', 'to', 'for', 'of',
            'jsx and props', 'jsx & props', 'jsx/props'  # Multi-word concepts, not standalone
            # Note: 'jsx' alone is valid as it's a React syntax extension
        }
        
        def is_core_technical_skill(term: str) -> bool:
            """
            Strict check if a term is a CORE technical skill (language, framework, tool, platform).
            Does NOT accept generic terms, verbs, or domain concepts.
            """
            if not term:
                return False
            
            term_lower = term.lower().strip()
            
            # Skip if it's a soft skill
            if is_soft_skill(term_lower):
                return False
            
            # Skip invalid standalone terms
            if term_lower in invalid_standalone_terms:
                return False
            
            # Skip if it's too short or too long
            if len(term_lower) < 2 or len(term_lower) > 50:
                return False
            
            # Skip common verbs and action words
            action_words = {'develop', 'create', 'build', 'design', 'implement', 'use', 'utilize',
                          'worked', 'working', 'work', 'developed', 'development', 'created',
                          'built', 'designed', 'implemented', 'using', 'used', 'utilized'}
            if term_lower in action_words:
                return False
            
            # Core technical skill patterns - only real technologies
            core_tech_patterns = [
                # Programming languages
                r'^(python|java|javascript|typescript|c\+\+|c#|go|rust|php|ruby|swift|kotlin|scala|r|'
                r'matlab|sql|html|css|sass|scss|less|dart|perl|shell|bash|powershell)$',
                
                # Frameworks and libraries
                r'^(react|angular|vue|ember|svelte|next\.js|nuxt\.js|gatsby|remix|'
                r'node\.js|express|koa|nest|django|flask|fastapi|bottle|pyramid|'
                r'spring|spring boot|hibernate|struts|play|quarkus|micronaut|'
                r'laravel|symfony|codeigniter|cakephp|rails|sinatra|grails|'
                r'\.net|asp\.net|core|entity framework|wpf|winforms|'
                r'redux|mobx|zustand|recoil|jotai|'
                r'react query|tanstack query|apollo|relay|urql|'
                r'tailwind css|bootstrap|material-ui|ant design|chakra ui|styled-components|'
                r'emotion|css-in-js|sass|less|stylus)$',
                
                # Databases
                r'^(mysql|postgresql|postgres|mongodb|redis|cassandra|couchdb|dynamodb|'
                r'oracle|sqlite|mariadb|neo4j|influxdb|elasticsearch|solr)$',
                
                # Tools and platforms
                r'^(git|github|gitlab|bitbucket|jira|confluence|trello|asana|'
                r'docker|kubernetes|k8s|jenkins|travis|circleci|github actions|gitlab ci|'
                r'aws|azure|gcp|heroku|vercel|netlify|firebase|supabase|'
                r'linux|unix|windows|macos|ios|android)$',
                
                # Protocols and standards
                r'^(http|https|rest|graphql|soap|grpc|websocket|tcp|udp|'
                r'oauth|jwt|ssl|tls|ssh|ftp|sftp)$',
                
                # File extensions (as standalone terms)
                r'^\.(js|ts|jsx|tsx|py|java|cpp|c|cs|php|rb|go|rs|swift|kt|scala|r|'
                r'sql|html|css|scss|sass|less|json|xml|yaml|yml|sh|bat|ps1)$'
            ]
            
            for pattern in core_tech_patterns:
                if re.match(pattern, term_lower, re.IGNORECASE):
                    return True
            
            # Check for known tech acronyms (2-5 uppercase letters)
            if re.match(r'^[A-Z]{2,5}$', term.strip()):
                # Common tech acronyms
                tech_acronyms = {'API', 'REST', 'SQL', 'HTML', 'CSS', 'JS', 'TS', 'JSX', 'TSX',
                               'AWS', 'GCP', 'CI', 'CD', 'UI', 'UX', 'IDE', 'CLI', 'SDK',
                               'HTTP', 'HTTPS', 'TCP', 'UDP', 'SSL', 'TLS', 'SSH', 'FTP',
                               'JWT', 'OAuth', 'GraphQL', 'gRPC', 'JSON', 'XML', 'YAML', 'DOM',
                               'BOM', 'AJAX', 'RPC', 'SOAP', 'REST', 'CRUD', 'ORM', 'MVC', 'MVP'}
                if term.strip() in tech_acronyms:
                    return True
            
            # Check for tech with version numbers (React 18, Python 3.9)
            if re.match(r'^[a-z]+\s*\d+', term_lower):
                base_term = re.sub(r'\s*\d+.*$', '', term_lower).strip()
                if is_core_technical_skill(base_term):
                    return True
            
            return False
        
        def is_domain_skill(term: str) -> bool:
            """
            Check if a term is a domain skill (like "Web Development", "API Integration").
            These are valid skills but should be separated from core technical skills.
            """
            if not term:
                return False
            
            term_lower = term.lower().strip()
            
            # Skip if it's a soft skill
            if is_soft_skill(term_lower):
                return False
            
            # Skip if it's too short or too long
            if len(term_lower) < 3 or len(term_lower) > 60:
                return False
            
            # Domain skill patterns
            domain_patterns = [
                r'\b(web development|frontend development|backend development|full stack development|'
                r'mobile development|ios development|android development|'
                r'api development|api integration|restful api|graphql api|'
                r'responsive design|web accessibility|ui/ux design|'
                r'component reusability|state management|'
                r'devops|ci/cd|cloud computing|microservices|serverless|'
                r'machine learning|data science|data analysis|nlp|computer vision)\b'
            ]
            
            for pattern in domain_patterns:
                if re.search(pattern, term_lower, re.IGNORECASE):
                    return True
            
            # Multi-word terms that contain tech words but are domain concepts
            words = term_lower.split()
            if len(words) >= 2 and len(words) <= 4:
                tech_domain_words = ['development', 'design', 'integration', 'management', 'architecture']
                tech_words = ['web', 'api', 'frontend', 'backend', 'mobile', 'responsive', 'component']
                if any(word in tech_domain_words for word in words) and any(word in tech_words for word in words):
                    return True
            
            return False
        
        def add_skill(term: str, is_domain: bool = False):
            """
            Add a skill with proper normalization and deduplication.
            """
            if not term or not term.strip():
                return
            
            term_clean = term.strip()
            normalized = normalize_skill(term_clean)
            
            if not normalized:
                return
            
            # Skip invalid standalone terms
            if normalized in invalid_standalone_terms:
                return
            
            # Get best form
            best_form = get_best_form(term_clean)
            
            if is_domain:
                domain_skills_set.add(best_form)
            else:
                # Store normalized -> best form mapping
                if normalized not in normalized_to_best_form:
                    normalized_to_best_form[normalized] = best_form
                else:
                    # If we already have this skill, prefer the better form
                    existing = normalized_to_best_form[normalized]
                    # Prefer forms with proper capitalization
                    if len(best_form) > len(existing) or (best_form[0].isupper() and not existing[0].isupper()):
                        normalized_to_best_form[normalized] = best_form
        
        # 1. Extract from "Technical Skills" section explicitly
        text_lines = text.split('\n')
        
        for i, line in enumerate(text_lines):
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            # Look for Technical Skills section header
            if re.match(r'^(technical\s+skills?|skills?|technical\s+expertise|technologies?|tech\s+stack)[:\s]*$', line_lower):
                # Extract skills from this section until next major section
                section_end_keywords = [
                    'education', 'experience', 'projects', 'certification', 'training',
                    'achievements', 'awards', 'languages', 'references', 'contact',
                    'summary', 'objective', 'profile', 'work experience', 'employment'
                ]
                
                # Extract content from this section
                for j in range(i + 1, len(text_lines)):
                    next_line = text_lines[j].strip()
                    if not next_line:
                        continue
                    
                    next_line_lower = next_line.lower()
                    # Check if we've hit a new major section
                    is_section_end = False
                    for end_keyword in section_end_keywords:
                        if re.match(rf'^{re.escape(end_keyword)}[:\s]*$', next_line_lower):
                            is_section_end = True
                            break
                    
                    if is_section_end:
                        break
                    
                    # Extract skills from this line
                    # First, try to extract multi-word technical terms (like "Tailwind CSS", "React Query")
                    multi_word_patterns = [
                        r'\b(Tailwind CSS|React Query|TanStack Query|React\.js|Node\.js|Next\.js|Nuxt\.js|'
                        r'TypeScript|JavaScript|React Native|Material-UI|Ant Design|Chakra UI|'
                        r'Styled Components|React Hooks|Redux Toolkit|GraphQL API|REST API|'
                        r'Web Development|Frontend Development|Backend Development|Full Stack|'
                        r'API Integration|Component Reusability|Responsive Design)\b'
                    ]
                    
                    for pattern in multi_word_patterns:
                        matches = re.finditer(pattern, next_line, re.IGNORECASE)
                        for match in matches:
                            term = match.group(0).strip()
                            if is_core_technical_skill(term):
                                add_skill(term, is_domain=False)
                            elif is_domain_skill(term):
                                add_skill(term, is_domain=True)
                    
                    # Then extract single-word and remaining terms
                    # Handle comma-separated, pipe-separated, colon-separated, semicolon-separated, or bullet-separated lists
                    skill_items = re.split(r'[,|;:•·▪▸▹▪▫◦‣⁃⁌⁍→➜➤○●\-]', next_line)
                    for item in skill_items:
                        item = item.strip()
                        if not item:
                            continue
                        
                        # Remove leading/trailing punctuation but preserve dots (for React.js, etc.)
                        item = re.sub(r'^[^\w\.]+|[^\w\.]+$', '', item)
                        
                        if not item:
                            continue
                        
                        # Skip if already extracted as part of multi-word term
                        # Check if this item is part of any multi-word term we already extracted
                        is_part_of_multiword = False
                        for extracted_term in list(normalized_to_best_form.values()) + list(domain_skills_set):
                            if item.lower() in extracted_term.lower() and item.lower() != extracted_term.lower():
                                is_part_of_multiword = True
                                break
                        
                        if is_part_of_multiword:
                            continue
                        
                        # Check if it's a core technical skill
                        if is_core_technical_skill(item):
                            add_skill(item, is_domain=False)
                        # Check if it's a domain skill
                        elif is_domain_skill(item):
                            add_skill(item, is_domain=True)
                        
                        # Also check for multi-word technical terms in the item itself
                        words = item.split()
                        if len(words) >= 2 and len(words) <= 4:
                            phrase = ' '.join(words)
                            if is_core_technical_skill(phrase):
                                add_skill(phrase, is_domain=False)
                            elif is_domain_skill(phrase):
                                add_skill(phrase, is_domain=True)
                break
        
        # 2. Extract from project tech stacks
        for project in projects:
            tech_stack = project.get("tech_stack", [])
            if isinstance(tech_stack, list):
                for tech in tech_stack:
                    if isinstance(tech, str):
                        tech_clean = tech.strip()
                        if is_core_technical_skill(tech_clean):
                            add_skill(tech_clean, is_domain=False)
                        elif is_domain_skill(tech_clean):
                            add_skill(tech_clean, is_domain=True)
            
            # Also check project summary for tech mentions
            summary = project.get("summary", "")
            if summary:
                # Remove HTML tags for better extraction
                summary_clean = re.sub(r'<[^>]+>', ' ', summary)
                
                # Extract technical terms from summary
                # Look for patterns like "using React, JavaScript, HTML" or "Tech Stack: React, JS"
                tech_mention_patterns = [
                    r'\b(?:using|with|built with|developed with|technologies?|tech stack|stack|tools?)[:\s]+([^<\.!?]+?)(?:<br>|\.|!|\?|$)',
                    r'tech\s+stack[:\s]+([^<\.!?]+?)(?:<br>|\.|!|\?|$)',
                    r'technologies?[:\s]+([^<\.!?]+?)(?:<br>|\.|!|\?|$)',
                ]
                
                for pattern in tech_mention_patterns:
                    tech_mentions = re.findall(pattern, summary_clean, re.IGNORECASE)
                    for mention in tech_mentions:
                        # Split by common separators
                        tech_items = re.split(r'[,|•·▪▸▹▪▫◦‣⁃⁌⁍→➜➤○●\-]', mention)
                        for item in tech_items:
                            item = item.strip()
                            if is_core_technical_skill(item):
                                add_skill(item, is_domain=False)
                            elif is_domain_skill(item):
                                add_skill(item, is_domain=True)
                
                # Extract standalone technical terms from summary
                # Look for known tech terms in the text
                known_tech_terms = [
                    r'\b(React|Angular|Vue|Node\.?js|JavaScript|TypeScript|HTML|CSS|'
                    r'Python|Java|C\+\+|C#|Go|Rust|PHP|Ruby|Swift|Kotlin|Scala|SQL|'
                    r'Django|Flask|FastAPI|Express|Spring|Laravel|Rails|\.NET|Next\.js|Nuxt\.js|'
                    r'MongoDB|PostgreSQL|MySQL|Redis|Docker|Kubernetes|AWS|Azure|GCP|'
                    r'Git|GitHub|GitLab|REST|GraphQL|Redux|Tailwind CSS|Bootstrap|'
                    r'React Query|TanStack Query)\b'
                ]
                
                for pattern in known_tech_terms:
                    matches = re.finditer(pattern, summary_clean, re.IGNORECASE)
                    for match in matches:
                        tech_term = match.group(0).strip()
                        if is_core_technical_skill(tech_term):
                            add_skill(tech_term, is_domain=False)
        
        # 3. Extract from technologies in keywords (if they came from resume)
        technologies = keywords.get("technologies", [])
        for tech in technologies:
            if isinstance(tech, str):
                tech_clean = tech.strip()
                if is_core_technical_skill(tech_clean):
                    add_skill(tech_clean, is_domain=False)
                elif is_domain_skill(tech_clean):
                    add_skill(tech_clean, is_domain=True)
        
        # 4. Extract from skills list (but filter to ensure they're technical)
        for skill in skills:
            if isinstance(skill, str):
                skill_clean = skill.strip()
                if is_core_technical_skill(skill_clean):
                    add_skill(skill_clean, is_domain=False)
                elif is_domain_skill(skill_clean):
                    add_skill(skill_clean, is_domain=True)
        
        # 5. Extract technical terms from project descriptions and responsibilities
        for project in projects:
            responsibilities = project.get("responsibilities", [])
            if isinstance(responsibilities, list):
                for resp in responsibilities:
                    if isinstance(resp, str):
                        # Remove HTML tags if present
                        resp_clean = re.sub(r'<[^>]+>', ' ', resp)
                        
                        # Extract core technical skills from responsibility text
                        # Look for known tech patterns
                        tech_patterns = [
                            r'\b(react|angular|vue|node\.?js|javascript|typescript|html|css|'
                            r'python|java|c\+\+|c#|go|rust|php|ruby|swift|kotlin|scala|sql|'
                            r'django|flask|fastapi|express|spring|laravel|rails|\.net|'
                            r'mongodb|postgresql|mysql|redis|docker|kubernetes|aws|azure|gcp|'
                            r'git|github|gitlab|rest|graphql|redux|tailwind css|bootstrap|'
                            r'react query|tanstack query|apollo|relay|material-ui|ant design|'
                            r'chakra ui|styled-components|sass|scss|less|webpack|vite|babel)\b'
                        ]
                        for pattern in tech_patterns:
                            matches = re.finditer(pattern, resp_clean, re.IGNORECASE)
                            for match in matches:
                                tech_term = match.group(0).strip()
                                if is_core_technical_skill(tech_term):
                                    add_skill(tech_term, is_domain=False)
                        
                        # Extract domain skills from responsibilities
                        domain_patterns = [
                            r'\b(api\s+integration|component\s+reusability|responsive\s+design|'
                            r'rest\s+api|graphql\s+api|web\s+development|frontend\s+development|'
                            r'backend\s+development|full\s+stack\s+development)\b'
                        ]
                        for pattern in domain_patterns:
                            matches = re.finditer(pattern, resp_clean, re.IGNORECASE)
                            for match in matches:
                                domain_term = match.group(0).strip()
                                if is_domain_skill(domain_term):
                                    add_skill(domain_term, is_domain=True)
        
        # 6. Extract from tools mentioned in projects
        for project in projects:
            tools = project.get("tools", [])
            if isinstance(tools, list):
                for tool in tools:
                    if isinstance(tool, str):
                        tool_clean = tool.strip()
                        if is_core_technical_skill(tool_clean):
                            add_skill(tool_clean, is_domain=False)
        
        # Get deduplicated core technical skills
        core_technical_skills = list(normalized_to_best_form.values())
        core_technical_skills.sort(key=lambda x: x.lower())
        
        # Get domain skills
        domain_skills = list(domain_skills_set)
        domain_skills.sort(key=lambda x: x.lower())
        
        # Return both lists
        return {
            "core_technical_skills": core_technical_skills[:25],  # Limit core skills
            "domain_skills": domain_skills[:10]  # Limit domain skills
        }
    
    def _extract_coding_topics(self, skills: List[str], keywords: Dict[str, List[str]], text_lower: str) -> List[str]:
        """
        Extract ONLY DSA/algorithm topics from resume. 
        DO NOT include frontend/UI topics (React, Components, Hooks, Routing, CSS, HTML, API Integration, etc.)
        These belong in Technical Interview section, NOT Coding section.
        """
        coding_topics = []
        
        # Normalization mapping for consistent output
        normalization_map = {
            'javascript': 'JavaScript',
            'js': 'JavaScript',
            'typescript': 'TypeScript',
            'ts': 'TypeScript',
            'python': 'Python',
            'java': 'Java',
            'c++': 'C++',
            'c#': 'C#',
            'sql': 'SQL',
            'sql queries': 'SQL Queries',
            'sql query': 'SQL Queries',
            'database query': 'SQL Queries',
            'arrays': 'Arrays',
            'strings': 'Strings',
            'arrays & strings': 'Arrays and Strings',
            'arrays and strings': 'Arrays and Strings',
            'hash table': 'Hash Tables',
            'hash tables': 'Hash Tables',
            'maps': 'Hash Tables',
            'linked list': 'Linked List',
            'stack': 'Stack',
            'queue': 'Queue',
            'recursion': 'Recursion',
            'two pointer': 'Two Pointers',
            'two pointers': 'Two Pointers',
            'sorting': 'Sorting',
            'searching': 'Searching',
            'graph': 'Basic Graph',
            'graphs': 'Basic Graph',
            'dynamic programming': 'Dynamic Programming',
            'time complexity': 'Time & Space Complexity',
            'space complexity': 'Time & Space Complexity',
            'complexity': 'Time & Space Complexity'
        }
        
        # Extract DSA topics ONLY if explicitly mentioned in resume
        dsa_keywords = {
            'arrays': 'Arrays',
            'strings': 'Strings',
            'hash table': 'Hash Tables',
            'hash tables': 'Hash Tables',
            'map': 'Hash Tables',
            'maps': 'Hash Tables',
            'linked list': 'Linked List',
            'stack': 'Stack',
            'queue': 'Queue',
            'recursion': 'Recursion',
            'two pointer': 'Two Pointers',
            'two pointers': 'Two Pointers',
            'sorting': 'Sorting',
            'searching': 'Searching',
            'graph': 'Basic Graph',
            'graphs': 'Basic Graph',
            'dynamic programming': 'Dynamic Programming',
            'time complexity': 'Time & Space Complexity',
            'space complexity': 'Time & Space Complexity',
            'complexity': 'Time & Space Complexity',
            'algorithm': 'Algorithms',
            'algorithms': 'Algorithms',
            'data structure': 'Data Structures',
            'data structures': 'Data Structures',
            'dsa': 'Data Structures'
        }
        
        # Check for DSA topics in resume text
        for keyword, topic_name in dsa_keywords.items():
            if keyword in text_lower:
                coding_topics.append(topic_name)
        
        # Extract JavaScript Logic (programming fundamentals) if JavaScript is mentioned
        # But NOT React, Components, Hooks, Routing, etc.
        if any('javascript' in s.lower() or 'js' in s.lower() for s in skills) or 'javascript' in text_lower:
            # Only add JavaScript Logic if JavaScript is used for programming/logic (not just UI)
            # Check if JavaScript is used for algorithms/logic, not just React/UI
            if any(term in text_lower for term in ['logic', 'algorithm', 'problem solving', 'programming', 'function', 'variable']):
                coding_topics.append('JavaScript Logic')
        
        # SQL Queries only if explicitly mentioned
        if any(term in text_lower for term in ['sql', 'sql query', 'database query', 'query', 'database']):
            coding_topics.append('SQL Queries')
        
        # Normalize topics
        normalized_topics = []
        for topic in coding_topics:
            topic_lower = topic.lower()
            # Check normalization map
            normalized = normalization_map.get(topic_lower, topic)
            # Handle "Arrays & Strings" vs "Arrays and Strings"
            if 'arrays' in topic_lower and 'strings' in topic_lower:
                normalized = 'Arrays and Strings'
            normalized_topics.append(normalized)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_topics = []
        for topic in normalized_topics:
            topic_lower = topic.lower()
            if topic_lower not in seen:
                seen.add(topic_lower)
                unique_topics.append(topic)
        
        return unique_topics
    
    def _extract_hr_skills(self, text_lower: str, keywords: Dict[str, List[str]]) -> List[str]:
        """Extract communication, teamwork, leadership skills for HR Interview"""
        hr_skills = []
        
        # Extract soft skills from text
        soft_skill_patterns = {
            'communication': ['communication', 'communicate', 'presentation', 'presenting'],
            'teamwork': ['team', 'collaboration', 'collaborate', 'teamwork', 'group'],
            'leadership': ['lead', 'leader', 'leadership', 'manage', 'management', 'mentor', 'supervise'],
            'problem solving': ['problem solving', 'problem-solve', 'analytical', 'analysis'],
            'adaptability': ['adapt', 'flexible', 'agile', 'scrum'],
            'time management': ['time management', 'deadline', 'prioritize']
        }
        
        for skill_name, patterns in soft_skill_patterns.items():
            if any(pattern in text_lower for pattern in patterns):
                hr_skills.append(skill_name.title())
        
        # Extract from job titles if available
        job_titles = keywords.get("job_titles", [])
        for title in job_titles:
            title_lower = title.lower()
            if 'lead' in title_lower or 'senior' in title_lower or 'manager' in title_lower:
                if 'Leadership' not in hr_skills:
                    hr_skills.append('Leadership')
        
        # Add strengths/weaknesses insights
        if len(hr_skills) >= 3:
            hr_skills.append('Strong interpersonal skills')
        if any(term in text_lower for term in ['project', 'deliver', 'achieve']):
            hr_skills.append('Project delivery experience')
        
        return list(dict.fromkeys(hr_skills))[:10]  # Remove duplicates, limit to 10
    
    def _extract_star_points(self, text: str, projects: List[Dict[str, str]], keywords: Dict[str, List[str]]) -> List[str]:
        """Extract STAR method points from resume for Behavioral Interview"""
        star_points = []
        
        # Extract from projects
        for project in projects:
            project_name = project.get("name", "")
            project_summary = project.get("summary", "")
            
            if project_name:
                star_points.append(f"Situation: Worked on {project_name}")
            
            if project_summary:
                # Try to extract action verbs
                action_verbs = ['developed', 'created', 'built', 'designed', 'implemented', 'optimized', 'improved']
                for verb in action_verbs:
                    if verb in project_summary.lower():
                        star_points.append(f"Action: {verb.capitalize()} solution for {project_name}")
                        break
        
        # Extract from text - look for achievement patterns
        achievement_patterns = [
            r'increased\s+[^.!?]{5,50}',
            r'reduced\s+[^.!?]{5,50}',
            r'improved\s+[^.!?]{5,50}',
            r'achieved\s+[^.!?]{5,50}',
            r'led\s+[^.!?]{5,50}',
            r'managed\s+[^.!?]{5,50}'
        ]
        
        for pattern in achievement_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                result_text = match.group(0).strip()
                if len(result_text) < 100:  # Keep it concise
                    star_points.append(f"Result: {result_text}")
        
        # Extract challenges/situations
        challenge_keywords = ['challenge', 'problem', 'issue', 'difficult', 'complex']
        sentences = text.split('.')
        for sentence in sentences[:20]:  # Check first 20 sentences
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in challenge_keywords) and len(sentence) > 20:
                star_points.append(f"Situation: {sentence.strip()[:80]}")
                break
        
        # Limit and ensure uniqueness
        return list(dict.fromkeys(star_points))[:8]  # Remove duplicates, limit to 8
    
    def _determine_coding_difficulty(self, experience_level: str, skills: List[str], 
                                    projects: List[Dict[str, str]], text_lower: str) -> str:
        """
        Calculate difficulty level based on resume depth and skill diversity.
        - Basic programming (HTML, CSS, JS) → Beginner
        - JavaScript + React + API usage → Easy to Medium
        - DSA terms explicitly mentioned → Medium to Hard
        """
        # Count programming languages detected
        programming_langs = ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'php', 'ruby', 'swift', 'kotlin', 'c', 'cpp']
        detected_langs = [lang for lang in programming_langs if lang in text_lower]
        
        # Check for basic programming only (HTML, CSS, basic JS)
        basic_programming = ['html', 'css']
        has_basic_only = any(tech in text_lower for tech in basic_programming) and len(detected_langs) <= 1
        
        # Check for JavaScript + React + API usage
        has_javascript = 'javascript' in text_lower or 'js' in text_lower
        has_react = 'react' in text_lower
        has_api = 'api' in text_lower or 'rest' in text_lower
        has_js_react_api = has_javascript and has_react and has_api
        
        # Check for DSA terms explicitly mentioned
        dsa_indicators = ['algorithm', 'data structure', 'dsa', 'data structures', 'arrays', 'strings', 
                         'hash table', 'linked list', 'stack', 'queue', 'recursion', 'two pointer',
                         'sorting', 'searching', 'graph', 'dynamic programming', 'complexity',
                         'leetcode', 'hackerrank', 'codechef', 'codeforces', 'competitive programming']
        has_dsa_mention = any(indicator in text_lower for indicator in dsa_indicators)
        
        # Count projects
        project_count = len(projects) if projects else 0
        
        # Calculate difficulty based on resume depth
        # Rule 1: Basic programming only → Beginner
        if has_basic_only and not has_javascript and not has_dsa_mention:
            return "Beginner"
        
        # Rule 2: JavaScript + React + API → Easy to Medium
        if has_js_react_api and not has_dsa_mention:
            return "Easy to Medium"
        
        # Rule 3: DSA terms explicitly mentioned → Medium to Hard
        if has_dsa_mention:
            # Adjust based on project count and skill diversity
            if project_count >= 3 and len(detected_langs) >= 2:
                return "Medium to Hard"
            else:
                return "Easy to Medium"
        
        # Rule 4: Multiple programming languages but no DSA → Easy to Medium
        if len(detected_langs) >= 2 and project_count >= 2:
            return "Easy to Medium"
        
        # Rule 5: Single language with projects → Beginner to Easy
        if len(detected_langs) == 1 and project_count >= 1:
            return "Beginner to Easy"
        
        # Default: Beginner
        return "Beginner"
    
    def _recommend_coding_platforms(self, skills: List[str], text_lower: str) -> List[str]:
        """
        Recommend coding platforms dynamically based on detected programming languages.
        Only show platforms if programming languages are detected in resume.
        """
        platforms = []
        
        # Detect programming languages in resume
        programming_langs = ['javascript', 'python', 'java', 'c++', 'c', 'cpp', 'c#', 'go', 'rust', 'php', 'ruby', 'swift', 'kotlin', 'typescript']
        detected_langs = [lang for lang in programming_langs if lang in text_lower]
        
        # Only recommend platforms if programming languages are detected
        if not detected_langs:
            return []  # Return empty if no programming languages found
        
        # Standard platforms to recommend if programming languages detected
        standard_platforms = ['LeetCode', 'HackerRank', 'CodeSignal', 'CodeChef']
        platforms.extend(standard_platforms)
        
        # Also check if platforms are explicitly mentioned in resume
        platform_mentions = {
            'leetcode': 'LeetCode',
            'hackerrank': 'HackerRank',
            'codesignal': 'CodeSignal',
            'codechef': 'CodeChef',
            'codeforces': 'Codeforces',
            'codewars': 'Codewars',
            'geeksforgeeks': 'GeeksforGeeks',
            'pramp': 'Pramp',
            'interviewbit': 'InterviewBit',
            'topcoder': 'TopCoder',
            'atcoder': 'AtCoder',
            'spoj': 'SPOJ'
        }
        
        # Add explicitly mentioned platforms
        for platform_key, platform_name in platform_mentions.items():
            if platform_key in text_lower and platform_name not in platforms:
                platforms.append(platform_name)
        
        # Remove duplicates and return
        return list(dict.fromkeys(platforms))
    
    def _recommend_coding_topics(self, skills: List[str], keywords: Dict[str, List[str]], 
                                text_lower: str) -> List[str]:
        """Recommend coding topics for practice"""
        topics = []
        
        # Core topics
        topics.extend(["Arrays & Strings", "Hash Tables", "Two Pointers", "Sliding Window"])
        
        # Based on skills
        if any('python' in s.lower() or 'java' in s.lower() for s in skills):
            topics.extend(["Dynamic Programming", "Greedy Algorithms"])
        
        if any('javascript' in s.lower() or 'react' in s.lower() for s in skills):
            topics.extend(["Tree Traversal", "Graph Algorithms"])
        
        # Based on domain
        if any(term in text_lower for term in ['database', 'sql', 'postgresql', 'mysql']):
            topics.append("SQL Queries")
        
        if any(term in text_lower for term in ['system design', 'architecture', 'microservices']):
            topics.extend(["System Design", "Scalability"])
        
        if any(term in text_lower for term in ['machine learning', 'ml', 'ai', 'data science']):
            topics.extend(["Algorithm Optimization", "Time Complexity"])
        
        return list(dict.fromkeys(topics))[:8]  # Remove duplicates, limit to 8
    
    def _generate_hr_evaluation_points(self) -> List[str]:
        """Generate common HR evaluation points"""
        return [
            "Communication skills and clarity of expression",
            "Cultural fit and alignment with company values",
            "Motivation and interest in the role",
            "Career goals and long-term aspirations",
            "Salary expectations and negotiation readiness",
            "Availability and notice period",
            "Team collaboration and interpersonal skills"
        ]
    
    def _generate_hr_suggestions(self, experience_level: str, skills: List[str], 
                                projects: List[Dict[str, str]], keywords: Dict[str, List[str]],
                                text_lower: str) -> List[str]:
        """Generate personalized HR interview suggestions"""
        suggestions = []
        
        # Experience-based suggestions
        if experience_level and experience_level not in ["Fresher", "Not specified", "Unknown"]:
            suggestions.append(f"Highlight your {experience_level} of experience and key achievements.")
        else:
            suggestions.append("Emphasize your learning ability, projects, and academic achievements.")
        
        # Skills-based suggestions
        if len(skills) >= 10:
            suggestions.append("Showcase your diverse technical skill set and adaptability.")
        elif len(skills) >= 5:
            suggestions.append("Focus on your core competencies and depth of knowledge.")
        
        # Project-based suggestions
        if len(projects) >= 3:
            suggestions.append("Prepare to discuss your projects in detail, focusing on impact and learnings.")
        elif len(projects) > 0:
            suggestions.append("Be ready to explain your project contributions and problem-solving approach.")
        
        # Domain-specific suggestions
        if any(term in text_lower for term in ['leadership', 'lead', 'mentor', 'team']):
            suggestions.append("Highlight your leadership experience and team collaboration skills.")
        
        if any(term in text_lower for term in ['startup', 'entrepreneur', 'founder']):
            suggestions.append("Emphasize your entrepreneurial mindset and ability to work in fast-paced environments.")
        
        # Default suggestions
        if not suggestions:
            suggestions.append("Prepare clear examples of your work and be ready to discuss your career goals.")
            suggestions.append("Research the company culture and align your answers with their values.")
        
        return suggestions[:4]  # Limit to 4 suggestions
    
    def _generate_star_guidance(self) -> Dict[str, str]:
        """Generate STAR method guidance"""
        return {
            "Situation": "Set the context: Describe the situation or challenge you faced. Be specific about when and where this occurred.",
            "Task": "Explain your responsibility: What was your role? What needed to be accomplished?",
            "Action": "Detail your actions: What specific steps did you take? Focus on your contributions, not the team's.",
            "Result": "Share the outcome: What was the result? Quantify achievements when possible (e.g., 'increased efficiency by 30%')."
        }
    
    def _generate_behavioral_tips(self, experience_level: str, projects: List[Dict[str, str]],
                                  keywords: Dict[str, List[str]], text_lower: str) -> List[str]:
        """Generate personalized behavioral interview tips"""
        tips = []
        
        # Experience-based tips
        if experience_level and experience_level not in ["Fresher", "Not specified", "Unknown"]:
            tips.append("Prepare 3-5 detailed STAR stories from your professional experience.")
        else:
            tips.append("Use academic projects, internships, or personal projects as STAR examples.")
        
        # Project-based tips
        if len(projects) >= 2:
            tips.append("Select your most impactful projects and prepare detailed STAR narratives for each.")
        
        # Common behavioral questions preparation
        tips.append("Practice answering: 'Tell me about a time you faced a challenge' and 'Describe a conflict you resolved'.")
        
        # Leadership/teamwork tips
        if any(term in text_lower for term in ['team', 'collaboration', 'group project']):
            tips.append("Prepare examples demonstrating teamwork, conflict resolution, and collaboration.")
        
        if any(term in text_lower for term in ['lead', 'manage', 'mentor', 'supervise']):
            tips.append("Have ready examples of leadership, decision-making, and mentoring others.")
        
        # Problem-solving tips
        if any(term in text_lower for term in ['problem', 'solve', 'debug', 'fix']):
            tips.append("Prepare stories showing your analytical thinking and problem-solving process.")
        
        # Default tips
        if len(tips) < 4:
            tips.append("Quantify your achievements with numbers, percentages, or time saved.")
            tips.append("Be specific and avoid vague answers. Use concrete examples.")
        
        return tips[:5]  # Limit to 5 tips
    
    def parse_resume(self, file_path: str, file_extension: str) -> Dict[str, Any]:
        """Parse resume and extract all relevant information"""
        try:
            # Extract text
            text = self.extract_text(file_path, file_extension)
            
            if not text or len(text.strip()) < 50:
                raise ValueError("Resume file appears to be empty or invalid")
            
            # Extract personal information
            name = self.extract_name(text)
            email = self.extract_email(text)
            
            # Extract skills
            skills = self.extract_skills(text)
            
            # Extract experience level
            experience_level = self.extract_experience_level(text)
            
            # Extract additional keywords
            keywords = self.extract_keywords(text)
            
            # Build parsed data
            parsed_data = {
                "name": name,
                "email": email,
                "skills": skills,
                "experience_level": experience_level,
                "keywords": keywords,
                "text_length": len(text),
                "extracted_text_preview": text[:500]  # First 500 chars for debugging
            }
            
            # Generate enhanced summary
            try:
                enhanced_summary = self.generate_enhanced_summary(parsed_data, text)
                parsed_data["summary"] = enhanced_summary
            except Exception as summary_error:
                # If summary generation fails, continue without it
                logger.warning(f"Failed to generate enhanced summary: {str(summary_error)}")
                parsed_data["summary"] = None
            
            return parsed_data
        except Exception as e:
            raise Exception(f"Error parsing resume: {str(e)}")

# Create global instance
resume_parser = ResumeParser()

